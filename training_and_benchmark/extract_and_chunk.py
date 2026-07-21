import os
import re
from pypdf import PdfReader
from docx import Document

def clean_and_extract_pdf(pdf_path):
    """
    Extracts text from NERC PDFs and strips repetitive running headers/footers.
    """
    reader = PdfReader(pdf_path)
    full_text = []
    
    print(f"Reading PDF: {os.path.basename(pdf_path)}...")
    
    for page in reader.pages:
        text = page.extract_text()
        if not text:
            continue
            
        # Filter out common running legal headers/footers line by line
        lines = text.split("\n")
        cleaned_lines = []
        for line in lines:
            if re.search(r"Standard CIP-\d{3}-\d|Page \d+ of \d+|Cyber Security|Reliability Standard", line, re.IGNORECASE):
                continue
            cleaned_lines.append(line)
            
        full_text.append("\n".join(cleaned_lines))
        
    return "\n".join(full_text)

def extract_docx_text(docx_path):
    """
    Extracts narratives from Word docs and flattens table questions into an audit text stream.
    """
    doc = Document(docx_path)
    full_text = []
    
    print(f"Reading Word Doc: {os.path.basename(docx_path)}...")
    
    # 1. Grab all regular paragraph text
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    # 2. Flatten evidence-request tables
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                # Deduplicate side-by-side text from merged cells
                unique_row_text = []
                for text in row_text:
                    if not unique_row_text or text != unique_row_text[-1]:
                        unique_row_text.append(text)
                full_text.append(" | ".join(unique_row_text))
                
    return "\n".join(full_text)

def chunk_by_requirement(raw_text):
    """
    Splits text by Requirement markers (R1, R2) to ensure clean thematic blocks.
    """
    # Split text wherever a major requirement boundary pattern occurs
    requirement_pattern = r"(?=\bRequirement\s+R\d+\b|\bR\d+\.\s+)"
    raw_chunks = re.split(requirement_pattern, raw_text)
    
    final_chunks = []
    for chunk in raw_chunks:
        chunk_stripped = chunk.strip()
        if not chunk_stripped:
            continue
            
        # Merge fragments under 400 characters into the previous chunk for context cohesion
        if len(chunk_stripped) < 400 and final_chunks:
            final_chunks[-1] += "\n\n" + chunk_stripped
        else:
            final_chunks.append(chunk_stripped)
            
    return final_chunks

def main():
    input_dir = "./compliance_inputs"
    output_dir = "./extracted_txt"

    os.makedirs(input_dir, exist_ok=True)
    os.makedirs(output_dir, exist_ok=True)

    # 1. Filter files upfront to ONLY process CIP-012 filenames
    all_files = [
        f for f in os.listdir(input_dir) 
        if not f.startswith('~$') and "CIP-012" in f.upper()
    ]

    if not all_files:
        print(f"No CIP-012 source files found inside '{input_dir}/'. Check file names.")
        return

    print(f"Targeted Recovery Mode: Identified {len(all_files)} CIP-012 files to chunk.")

    for file_name in all_files:
        file_path = os.path.join(input_dir, file_name)
        base_name = os.path.splitext(file_name)[0].upper()

        if file_name.lower().endswith('.pdf'):
            raw_text = clean_and_extract_pdf(file_path)
        elif file_name.lower().endswith('.docx'):
            raw_text = extract_docx_text(file_path)
        else:
            print(f"Skipping unsupported file type: {file_name}")
            continue

        raw_text = re.sub(r'\n\s*\n', '\n\n', raw_text)
        chunks = chunk_by_requirement(raw_text)

        output_file_path = os.path.join(output_dir, f"{base_name}_chunks.txt")
        with open(output_file_path, "w", encoding="utf-8") as f:
            for idx, chunk in enumerate(chunks):
                f.write(f"--- CHUNK {idx + 1} ---\n")
                f.write(chunk)
                f.write("\n\n")

        print(f"Successfully recovered {len(chunks)} chunks -> {output_file_path}\n")

    print("Targeted run complete. Your CIP-012 chunks have been injected back into 'extracted_txt/'.")

if __name__ == "__main__":
    main()
